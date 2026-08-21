from __future__ import annotations

from io import BytesIO
from typing import Callable

from docx import Document as DocxDocument
from docx.shared import Pt

from ocr_extraction import DocumentResult, PageResult, TextBox

LOW_CONFIDENCE_MARK_LEFT = "【"
LOW_CONFIDENCE_MARK_RIGHT = "】"

FlaggedItemLimit = 20


def _lines_of(column: list[TextBox], conf_threshold: float) -> list[str]:
    lines = []
    for box in column:
        text = box.text
        if box.confidence < conf_threshold:
            text = f"{LOW_CONFIDENCE_MARK_LEFT}{text}{LOW_CONFIDENCE_MARK_RIGHT}"
        lines.append(text)
    return lines


def _page_flag_summary(page: PageResult) -> str | None:
    if not page.flagged:
        return None
    items = ", ".join(f'"{box.text}" ({box.confidence:.0%})' for box in page.flagged[:FlaggedItemLimit])
    remainder = "" if len(page.flagged) <= FlaggedItemLimit else f" (+{len(page.flagged) - FlaggedItemLimit} more)"
    return f"Tekshirish tavsiya etiladi / Low-confidence: {items}{remainder}"


def to_text(doc: DocumentResult, conf_threshold: float = 0.80) -> str:
    parts: list[str] = []
    for page in doc.pages:
        parts.append(f"===== Page {page.page_number} ({page.source}, confidence {page.mean_confidence:.0%}) =====")
        for column_index, column in enumerate(page.columns):
            if len(page.columns) > 1:
                parts.append(f"--- column {column_index + 1} ---")
            parts.extend(_lines_of(column, conf_threshold))
        flag_summary = _page_flag_summary(page)
        if flag_summary:
            parts.append(flag_summary)
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def to_markdown(doc: DocumentResult, conf_threshold: float = 0.80) -> str:
    parts: list[str] = [f"# {doc.source_name}", ""]
    for page in doc.pages:
        parts.append(f"## Page {page.page_number}")
        parts.append(f"*source: {page.source}, mean confidence: {page.mean_confidence:.0%}*")
        parts.append("")
        if len(page.columns) > 1:
            for column_index, column in enumerate(page.columns):
                parts.append(f"### Column {column_index + 1}")
                parts.append("")
                parts.extend(_lines_of(column, conf_threshold))
                parts.append("")
        else:
            for column in page.columns:
                parts.extend(_lines_of(column, conf_threshold))
                parts.append("")
        flag_summary = _page_flag_summary(page)
        if flag_summary:
            parts.append(f"> ⚠️ {flag_summary}")
            parts.append("")
    return "\n".join(parts).strip() + "\n"


def to_docx_bytes(doc: DocumentResult, conf_threshold: float = 0.80) -> bytes:
    document = DocxDocument()
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading(doc.source_name, level=0)

    for page in doc.pages:
        document.add_heading(f"Page {page.page_number}", level=1)
        document.add_paragraph(f"source: {page.source} | mean confidence: {page.mean_confidence:.0%}")

        if len(page.columns) > 1:
            for column_index, column in enumerate(page.columns):
                document.add_heading(f"Column {column_index + 1}", level=2)
                for line in _lines_of(column, conf_threshold):
                    document.add_paragraph(line)
        else:
            for column in page.columns:
                for line in _lines_of(column, conf_threshold):
                    document.add_paragraph(line)

        flag_summary = _page_flag_summary(page)
        if flag_summary:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"⚠ {flag_summary}")
            run.italic = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


Formatter = Callable[..., str | bytes]

FORMATTERS: dict[str, Formatter] = {
    "text": to_text,
    "txt": to_text,
    "markdown": to_markdown,
    "md": to_markdown,
    "word": to_docx_bytes,
    "docx": to_docx_bytes,
}

FORMAT_EXTENSIONS: dict[str, str] = {
    "text": "txt", "txt": "txt",
    "markdown": "md", "md": "md",
    "word": "docx", "docx": "docx",
}
